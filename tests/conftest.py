from __future__ import annotations

from pathlib import Path
from typing import Any
from types import SimpleNamespace

import fitz
import pytest
from docx import Document
from ebooklib import epub

import sys

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from litreel import create_app
from litreel.services.gemini_runner import BookConcepts, SlideConcept

TEST_PASSWORD = "Testpass123!"


class DummyGemini:
    def __init__(self) -> None:
        self.called_with: Path | None = None
        self.text_payload: str | None = None
        self.chunk_calls: list[dict[str, Any]] = []
        self.document_parser = self._parse_document

    def _parse_document(self, document_path: Path | str) -> str:
        self.called_with = Path(document_path)
        return "Dummy extracted text for testing."

    def pdf_parser(self, pdf_path: Path | str) -> str:
        return self._parse_document(pdf_path)

    def generate_from_pdf(self, pdf_path: Path | str) -> BookConcepts:
        _ = pdf_path
        return self.generate_from_text("Dummy extracted text for testing.")

    def generate_from_text(self, text: str) -> BookConcepts:
        self.text_payload = text
        return BookConcepts(
            concepts=[
                SlideConcept(
                    name="Hidden Story",
                    description="A shocking untold history.",
                    slides=[
                        "Hook slide",
                        "Middle beat",
                        "Twist ending",
                    ],
                )
            ]
        )

    def generate_from_chunks(
        self,
        *,
        chunks,
        reference_concept=None,
        user_context=None,
    ) -> BookConcepts:
        self.chunk_calls.append(
            {
                "chunks": list(chunks),
                "reference": reference_concept,
                "context": user_context,
            }
        )
        return BookConcepts(
            concepts=[
                SlideConcept(
                    name="Contextual Hook",
                    description="A remix rooted in retrieved passages.",
                    slides=["New hook", "Supporting beat", "Final punch"],
                )
            ]
        )


class DummyRagService:
    def __init__(self) -> None:
        self.is_enabled = False
        self.can_background_ingest = True
        self.book_id = "sb-001"
        self.return_chunks = ["Chunk A", "Chunk B"]
        self.random_chunks = ["Random Chunk A", "Random Chunk B", "Random Chunk C"]
        self.ingest_calls: list[dict[str, Any]] = []
        self.retrieve_calls: list[dict[str, Any]] = []
        self.delete_calls: list[str] = []
        self.random_calls: list[dict[str, Any]] = []

    def ingest_book(self, *, title: str, text: str) -> str | None:
        self.ingest_calls.append({"title": title, "text": text})
        if not self.is_enabled:
            return None
        return self.book_id

    def get_relevant_chunks(self, book_id: str, query_text: str, match_count: int | None = None):
        self.retrieve_calls.append({"book_id": book_id, "query": query_text, "count": match_count})
        return list(self.return_chunks)

    def delete_book(self, book_id: str) -> None:
        self.delete_calls.append(book_id)

    def sample_random_chunks(self, book_id: str, sample_size: int = 75) -> list[str]:
        self.random_calls.append({"book_id": book_id, "sample_size": sample_size})
        return list(self.random_chunks[:sample_size])


class DummyArousalClient:
    def __init__(self) -> None:
        self.is_ready = True
        self.rankings: list[Any] = []
        self.calls: list[list[str]] = []

    def score_chunks(self, chunks: list[str]):
        self.calls.append(list(chunks))
        if self.rankings:
            return self.rankings
        return [
            SimpleNamespace(text=chunk, score=float(idx))
            for idx, chunk in enumerate(chunks)
        ]


class DummyStock:
    def __init__(self) -> None:
        self.last_query: str | None = None

    def search(self, query: str) -> list[dict[str, Any]]:
        self.last_query = query
        return [
            {
                "id": "stock-1",
                "url": "https://example.com/large.jpg",
                "thumbnail": "https://example.com/thumb.jpg",
                "photographer": "Test",
            }
        ]


class DummyRenderer:
    def __init__(self, root: Path) -> None:
        self.root = root
        self.called_with: tuple[int, int | None, str | None] | None = None

    def render_project(
        self,
        project,
        concept_id: int | None = None,
        voice: str | None = None,
        *,
        warnings: list[str] | None = None,
    ) -> Path:
        self.called_with = (project.id, concept_id, voice)
        target = self.root / f"project_{project.id}.mp4"
        target.write_bytes(b"fake")
        return target


@pytest.fixture
def sample_pdf(tmp_path: Path) -> Path:
    pdf_path = tmp_path / "sample.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "This is a viral-ready nonfiction excerpt.")
    doc.save(pdf_path)
    doc.close()
    return pdf_path


@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    doc = Document()
    doc.add_heading("Docx Sample", level=1)
    doc.add_paragraph("This DOCX file shares the same viral-ready nonfiction excerpt.")
    docx_path = tmp_path / "sample.docx"
    doc.save(docx_path)
    return docx_path


@pytest.fixture
def sample_epub(tmp_path: Path) -> Path:
    book = epub.EpubBook()
    book.set_identifier("sample-book")
    book.set_title("Sample EPUB")
    book.set_language("en")

    chapter = epub.EpubHtml(title="Intro", file_name="intro.xhtml", lang="en")
    chapter.content = (
        "<h1>Intro</h1><p>This EPUB holds the same viral-ready nonfiction excerpt.</p>"
    )
    book.add_item(chapter)
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    book.spine = ["nav", chapter]
    book.toc = (epub.Link("intro.xhtml", "Intro", "intro"),)

    epub_path = tmp_path / "sample.epub"
    epub.write_epub(str(epub_path), book)
    return epub_path


@pytest.fixture
def app(tmp_path: Path):
    uploads = tmp_path / "uploads"
    uploads.mkdir()
    gemini = DummyGemini()
    stock = DummyStock()
    renderer = DummyRenderer(tmp_path)
    rag = DummyRagService()
    arousal = DummyArousalClient()

    app = create_app(
        {
            "TESTING": True,
            "SQLALCHEMY_DATABASE_URI": f"sqlite:///{tmp_path / 'test.db'}",
            "UPLOAD_FOLDER": str(uploads),
            "GEMINI_SERVICE": gemini,
            "STOCK_IMAGE_SERVICE": stock,
            "VIDEO_RENDERER": renderer,
            "RAG_SERVICE": rag,
            "AROUSAL_CLIENT": arousal,
            "SUPABASE_URL": "",
            "SUPABASE_API_KEY": "",
        }
    )

    app.config["_dummy_services"] = {
        "gemini": gemini,
        "stock": stock,
        "renderer": renderer,
        "rag": rag,
        "arousal": arousal,
    }
    yield app


@pytest.fixture
def auth_client_factory(app):
    def _factory(email: str = "founder@example.com", password: str = TEST_PASSWORD):
        client = app.test_client()
        response = client.post(
            "/api/auth/signup",
            json={"email": email, "password": password},
        )
        assert response.status_code == 201
        return client, {"email": email, "password": password}

    return _factory


@pytest.fixture
def client(auth_client_factory):
    client, _ = auth_client_factory()
    return client


@pytest.fixture
def dummy_services(app):
    return app.config["_dummy_services"]
