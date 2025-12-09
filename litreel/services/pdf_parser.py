from __future__ import annotations

"""Document text extraction helpers."""

from pathlib import Path
import re
from typing import Iterable

import fitz


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".epub"}


def extract_text_from_document(document_path: Path | str) -> str:
    """Return normalized text content from a supported document type."""

    path = Path(document_path)
    if not path.exists():
        raise FileNotFoundError(f"Document not found: {path}")

    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_text_from_pdf(path)
    if suffix == ".docx":
        return extract_text_from_docx(path)
    if suffix == ".epub":
        return extract_text_from_epub(path)
    raise ValueError(
        f"Unsupported file extension '{suffix}'. Supported extensions: {', '.join(sorted(SUPPORTED_EXTENSIONS))}."
    )


def extract_text_from_pdf(pdf_path: Path | str) -> str:
    path = Path(pdf_path)
    if not path.exists():
        raise FileNotFoundError(f"PDF not found: {path}")

    doc = fitz.open(path)
    try:
        parts: list[str] = []
        for page_index in range(doc.page_count):
            page_text = doc.load_page(page_index).get_text("text").strip()
            if page_text:
                parts.append(page_text)
        return _normalize_text(parts)
    finally:
        doc.close()


def extract_text_from_docx(docx_path: Path | str) -> str:
    try:
        from docx import Document  # type: ignore[import]
    except ImportError as exc:  # pragma: no cover - environment safeguard
        raise RuntimeError("DOCX extraction requires the 'python-docx' package.") from exc

    path = Path(docx_path)
    if not path.exists():
        raise FileNotFoundError(f"DOCX not found: {path}")

    document = Document(path)
    blocks: list[str] = []

    def _append_text(candidate: str | None) -> None:
        if candidate:
            stripped = candidate.strip()
            if stripped:
                blocks.append(stripped)

    for paragraph in document.paragraphs:
        _append_text(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                _append_text(cell.text)

    return _normalize_text(blocks)


def extract_text_from_epub(epub_path: Path | str) -> str:
    try:
        from ebooklib import epub, ITEM_DOCUMENT  # type: ignore[import]
        from bs4 import BeautifulSoup  # type: ignore[import]
    except ImportError as exc:  # pragma: no cover - environment safeguard
        raise RuntimeError(
            "EPUB extraction requires the 'ebooklib' and 'beautifulsoup4' packages."
        ) from exc

    path = Path(epub_path)
    if not path.exists():
        raise FileNotFoundError(f"EPUB not found: {path}")

    book = epub.read_epub(path)
    segments: list[str] = []
    for item in book.get_items_of_type(ITEM_DOCUMENT):
        soup = BeautifulSoup(item.get_body_content(), "html.parser")
        text = soup.get_text(separator="\n", strip=True)
        if text:
            segments.append(text)

    return _normalize_text(segments)


def _normalize_text(chunks: Iterable[str]) -> str:
    combined = "\n\n".join(chunks)
    combined = combined.replace("\r\n", "\n").replace("\r", "\n")
    combined = re.sub(r"[ \t]+\n", "\n", combined)
    combined = re.sub(r"\n{3,}", "\n\n", combined)
    return combined.strip()
